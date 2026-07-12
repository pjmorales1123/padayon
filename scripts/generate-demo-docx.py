from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

out_path = r"C:\Users\Prince\Documents\PADAYON Demo Materials\PADAYON Live Demo Script.docx"

def set_cell_shading(cell, color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)

def bold_run(paragraph, text):
    run = paragraph.add_run(text)
    run.bold = True
    return run

doc = Document()

# Title
title = doc.add_heading('PADAYON Live Demo Script', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.runs[0]
run.font.size = Pt(24)
run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.add_run('3-minute live demo + 1:30 for methods | https://padayon-theta.vercel.app').italic = True

doc.add_paragraph()

# Pre-demo checklist
doc.add_heading('Pre-demo checklist (do this 5 minutes before judges arrive)', level=1)
items = [
    'Open terminal in the padayon repo and run:  node scripts/gemma4-scale.js up',
    'Wait until status shows replicas: 1 (check with: node scripts/gemma4-scale.js status)',
    'Seed demo personas: POST https://padayon-theta.vercel.app/api/seed-personas',
    'Open two browser tabs:',
    '   • Gemma 4 demo: https://padayon-theta.vercel.app/demo?userId=demo-bisaya-learner&model=gemma-4',
    '   • Fallback demo: https://padayon-theta.vercel.app/demo?userId=demo-struggling-student&model=fallback',
    'Have these files ready on screen (also in this folder):',
    '   • IMG_20260629_080203.jpg (handwritten photosynthesis notes)',
    '   • raisin_in_the_sun_lesson (1).pdf',
]
for item in items:
    p = doc.add_paragraph(item, style='List Bullet')
    p.paragraph_format.space_after = Pt(4)

# Materials section
doc.add_heading('Demo materials included in this folder', level=1)
doc.add_paragraph('6 note photos (handwritten photosynthesis notes) and 1 PDF (Raisin in the Sun lesson).')

# Demo flow table
doc.add_heading('Demo flow: 3 minutes 30 seconds', level=1)
table = doc.add_table(rows=1, cols=5)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
headers = ['Time', 'Step', 'What you say', 'What you do', 'Expected result']
for i, text in enumerate(headers):
    hdr_cells[i].text = text
    set_cell_shading(hdr_cells[i], 'D9E2F3')
    for paragraph in hdr_cells[i].paragraphs:
        for run in paragraph.runs:
            run.font.bold = True

rows = [
    (
        '0:00–0:30',
        'Hook: messy notes → study pack',
        '“Students come to class with messy notes. PADAYON turns those into a school-aligned study pack in seconds.”',
        'On Gemma 4 tab, click 📎/📷 and upload IMG_20260629_080203.jpg.',
        'Chat shows: “Got your image notes… I created Clean Notes, Reviewer, Flashcards, Quiz, Summary… Does this match what your teacher discussed?” Badge: Gemma 4 · AMD/Fireworks.'
    ),
    (
        '0:30–1:15',
        'Personalized explanation',
        '“Same lesson, different students. Juan is Cebuano-first. Watch how PADAYON adapts language while keeping the academic term.”',
        'Type: Explain photosynthesis',
        'Reply is bilingual: Cebuano-first, then English, includes Photosynthesis, asks one guiding question. No leaked thinking.'
    ),
    (
        '1:15–2:00',
        'Materials on demand',
        '“Now Juan wants to review. PADAYON retrieves what it already built.”',
        'Type: Show my flashcards',
        'Interactive flashcard widget appears. First card: “What is Photosynthesis?” Tap to flip.'
    ),
    (
        '2:00–2:30',
        'Truthful model badge',
        '“PADAYON never fakes the model. If Gemma is down, it falls back transparently. Let me switch to Fallback.”',
        'Switch selector to Fallback · DeepSeek V4 Flash, type: I don\'t get quadratic equations. It looks hard.',
        'Badge flips to DeepSeek V4 Flash · Fireworks. Reply is warm, step-by-step, no meta.'
    ),
    (
        '2:30–3:00',
        'Back to Gemma + school-aligned scope',
        '“Switch back to Gemma, and PADAYON stays bounded to the class lesson. If a student asks beyond the lesson, it asks first.”',
        'Switch back to Gemma 4, type: What is the Calvin cycle?',
        'Reply: “The Calvin cycle is related to photosynthesis, but it may be beyond your current class lesson. Focus first on…”'
    ),
    (
        '3:00–3:30',
        'Wrap / mastery map',
        '“Everything lives in a growing topic workspace: uploads, notes, quiz progress, and a mastery map tied to the teacher’s lesson.”',
        'Click “Open your study pack” link or open Library.',
        'Topic page shows Photosynthesis with tabs: Overview, Original Notes, Clean Notes, Reviewer, Flashcards, Quiz, Story, Progress.'
    ),
]

for r in rows:
    row_cells = table.add_row().cells
    for i, text in enumerate(r):
        row_cells[i].text = text

doc.add_paragraph()

# Exact prompts
doc.add_heading('Exact prompts to use', level=1)
prompts = [
    'Upload: (no text — just upload IMG_20260629_080203.jpg)',
    'Explain photosynthesis',
    'Show my flashcards',
    'I don\'t get quadratic equations. It looks hard. (on Fallback tab)',
    'What is the Calvin cycle? (back on Gemma tab)',
]
for p in prompts:
    doc.add_paragraph(p, style='List Number')

# Narrator script
doc.add_heading('Narrator script (speak naturally)', level=1)
script = (
    "Students don’t need another chatbot. They need help understanding what their teacher already taught.\n\n"
    "Here is Juan’s messy photo of notes. PADAYON reads it, builds a study pack, and asks: does this match your teacher’s lesson? "
    "That keeps the AI inside the class scope.\n\n"
    "Same lesson, different student. Juan is Cebuano-first, so PADAYON explains in Cebuano first, then English, and always keeps the academic term.\n\n"
    "Need to review? Flashcards are already built and ready.\n\n"
    "PADAYON is honest about the model. I can switch to Fallback — still clean, no leaked thinking — then back to Gemma.\n\n"
    "Finally, everything becomes a growing topic workspace: uploads, notes, quizzes, and a mastery map."
)
doc.add_paragraph(script)

# Fallback plan
doc.add_heading('Fallback plan', level=1)
fb = [
    'If Gemma badge shows Fallback · Fireworks on first try, wait 30 seconds and retry — the replica is still warming.',
    'If anything breaks, stay on Fallback · DeepSeek V4 Flash — it already produces clean, deterministic replies.',
    'After the demo, scale Gemma down to avoid cost: node scripts/gemma4-scale.js down',
]
for item in fb:
    doc.add_paragraph(item, style='List Bullet')

# Save
doc.save(out_path)
print('Created:', out_path)
